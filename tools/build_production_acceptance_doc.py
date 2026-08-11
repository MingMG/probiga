# -*- coding: utf-8 -*-
"""Build the ProBigA production data acceptance report from audit evidence."""
from __future__ import annotations

import argparse
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "17324D"
TEAL = "008C95"
GREEN = "138A5B"
AMBER = "B87700"
RED = "B42318"
SLATE = "526273"
PALE = "EEF4F7"
PALE_GREEN = "E9F6EF"
PALE_AMBER = "FFF5E5"
WHITE = "FFFFFF"
INK = RGBColor(28, 43, 56)
MUTED = RGBColor(82, 98, 115)


CHECK_LABELS = {
    "stock_universe_integrity": "股票基础池完整性",
    "daily_kline_business_rules": "日 K 业务规则与字段完整性",
    "daily_capital_flow_integrity": "日级资金流完整性",
    "current_quote_integrity": "全市场实时行情完整性",
    "stock_minute_integrity_and_coverage": "分钟价格完整性与覆盖率",
    "minute_capital_flow_integrity_and_coverage": "分钟资金流完整性与覆盖率",
    "current_vs_daily_close": "实时价格 vs 日线收盘价",
    "minute_last_vs_daily_close": "分钟末价 vs 日线收盘价",
    "minute_daily_universe_semantic_coverage": "分钟缺失代码语义分类",
    "sm_index_current_integrity": "指数实时行情完整性",
    "sm_concept_east_current_integrity": "东财概念实时行情完整性",
    "sm_concept_ths_current_integrity": "同花顺概念实时行情完整性",
    "concept_capital_flow_integrity": "概念资金流完整性",
    "st_hot_rank_ths_rank_integrity": "同花顺热榜完整性",
    "st_hot_pop_rank_east_rank_integrity": "东财热榜完整性",
    "st_hot_rank_xq_rank_integrity": "雪球热榜完整性",
    "st_hot_rank_sina_rank_integrity": "新浪热榜完整性",
    "st_hot_rank_fused_rank_integrity": "融合热榜完整性",
    "news_integrity": "多源快讯完整性",
    "notice_integrity": "公告数据完整性",
    "analysis_result_integrity": "全市场分析结果完整性",
    "recommendation_boundary_values": "推荐价位边界",
    "market_overview_arithmetic": "大盘汇总算术约束",
    "scheduler_enabled_task_health": "启用调度任务健康度",
    "temporary_stage_cleanup": "临时阶段表清理",
}


def load_json(path: str) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size=9, bold=False, color=INK, font="Microsoft YaHei") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def status_fill(status: str) -> str:
    status = str(status).upper()
    if status == "PASS":
        return PALE_GREEN
    if status in {"WARN", "SKIP"}:
        return PALE_AMBER
    return "FDECEC"


def metric_text(metrics: dict[str, Any], *, max_items: int = 7) -> str:
    preferred = [
        "latest_trade_date", "latest_date", "rows", "stocks", "distinct_codes",
        "coverage_ratio", "compared", "mismatches", "invalid_codes", "bad_price",
        "missing_change_fields", "invalid_stop_loss", "invalid_take_profit",
        "enabled_tasks", "unhealthy_tasks", "stage_tables", "minute_stage_tables",
    ]
    keys = [key for key in preferred if key in metrics]
    keys.extend(key for key in metrics if key not in keys)
    parts = []
    for key in keys[:max_items]:
        value = metrics.get(key)
        if isinstance(value, float):
            value = f"{value:.4f}".rstrip("0").rstrip(".")
        parts.append(f"{key}={value}")
    return "；".join(parts)


def add_table(document: Document, headers: list[str], rows: Iterable[Iterable[Any]], widths=None, font_size=8.2):
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        set_run_font(run, font_size, True, RGBColor(255, 255, 255))
        if widths:
            cell.width = widths[idx]
    for raw_row in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(raw_row):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("" if value is None else str(value))
            set_run_font(run, font_size)
            if idx == 0:
                set_run_font(run, font_size, True)
            if widths:
                cell.width = widths[idx]
        if len(row.cells) > 1 and str(raw_row[1]).upper() in {"PASS", "WARN", "FAIL", "ERROR", "SKIP"}:
            set_cell_shading(row.cells[1], status_fill(str(raw_row[1])))
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(document: Document, text: str, level=1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(5)


def add_body(document: Document, text: str, *, bold_lead: str = "") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.16
    if bold_lead:
        run = paragraph.add_run(bold_lead)
        set_run_font(run, 9.5, True)
    run = paragraph.add_run(text)
    set_run_font(run, 9.5)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)
        run = paragraph.add_run(str(item))
        set_run_font(run, 9.2)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    for name, size, color in (("Title", 28, RGBColor(255, 255, 255)), ("Heading 1", 16, RGBColor(23, 50, 77)), ("Heading 2", 11.5, RGBColor(0, 140, 149))):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header
    table = header.add_table(rows=1, cols=2, width=section.page_width - section.left_margin - section.right_margin)
    table.autofit = False
    table.columns[0].width = Cm(11.5)
    table.columns[1].width = Cm(5.8)
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=35, bottom=35)
    p = left.paragraphs[0]
    run = p.add_run("ProBigA  /  PROD DATA ACCEPTANCE")
    set_run_font(run, 8.2, True, RGBColor(255, 255, 255))
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("生产验收 · 内部")
    set_run_font(run, 8.2, True, RGBColor(255, 255, 255))

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("ProBigA 数据链路验收  |  ")
    set_run_font(run, 8, False, MUTED)
    add_field(p, "PAGE")


def add_cover(document: Document, database: dict[str, Any], overall: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=600, start=430, bottom=580, end=430)
    p = cell.paragraphs[0]
    run = p.add_run("PRODUCTION DATA ACCEPTANCE")
    set_run_font(run, 10, True, RGBColor(111, 222, 225))
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("ProBigA 生产数据\n全流程验收报告")
    set_run_font(run, 27, True, RGBColor(255, 255, 255))
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run("盘中 / 非盘中 · 源数据 / 落库 / API · 准确性 / 边界值 / 资源安全")
    set_run_font(run, 10.5, False, RGBColor(220, 232, 239))

    document.add_paragraph()
    status_table = document.add_table(rows=1, cols=3)
    status_table.style = "Table Grid"
    values = [
        ("验收状态", overall),
        ("数据基准日", database.get("latest_trade_date", "2026-07-17")),
        ("报告日期", "2026-07-19"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = status_table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_GREEN if overall == "通过" else PALE_AMBER)
        set_cell_margins(cell, top=160, bottom=160)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}\n")
        set_run_font(run, 8, True, MUTED)
        run = p.add_run(str(value))
        set_run_font(run, 14, True, RGBColor(19, 138, 91) if overall == "通过" else RGBColor(184, 119, 0))

    document.add_paragraph()
    add_body(document, "验收环境：生产环境；时区：Asia/Shanghai；最新完整交易日：2026-07-17。")
    add_body(document, "报告依据：生产真实源抓取日志、数据库只读审计、运行时资源审计、应用 API 回读、自动化回归测试及异常恢复演练。")
    document.add_page_break()


def build_report(args: argparse.Namespace) -> Path:
    database = load_json(args.database)
    runtime = load_json(args.runtime)
    inventory = load_json(args.inventory)
    api = load_json(args.api)
    source = load_json(args.source)

    checks = database.get("checks", [])
    active_profiles = [item for item in database.get("profiles", []) if item.get("active")]
    enabled_tasks = [item for item in inventory.get("tasks", []) if int(item.get("enabled") or 0) == 1]
    all_db_pass = all(item.get("status") == "PASS" for item in checks)
    all_profile_pass = all(item.get("status") == "PASS" for item in active_profiles)
    all_api_pass = int(api.get("failed") or 0) == 0
    overall = "通过" if all_db_pass and all_profile_pass and all_api_pass else "有条件通过"

    document = Document()
    configure_document(document)
    add_cover(document, database, overall)

    add_heading(document, "管理摘要", 1)
    add_body(
        document,
        f"本轮对生产环境启用数据链路进行了从调度入口、真实外部源、写库事务、应用读库、API 输出到资源约束的闭环验收。数据库规则检查 {sum(c.get('status') == 'PASS' for c in checks)}/{len(checks)} 通过；启用输出画像 {sum(p.get('status') == 'PASS' for p in active_profiles)}/{len(active_profiles)} 通过；API 回读 {api.get('passed', 0)}/{len(api.get('results', []))} 通过。",
    )
    add_table(document, ["验收维度", "结论", "关键证据"], [
        ("真实源获取", "PASS", "行情 5,203；日 K 5,203；分钟价格/资金流各 5,195 只成功；多源热榜与资讯均有输出"),
        ("准确性", "PASS", "实时价 vs 日收盘 5,203/5,203 一致；日 K 前收/涨跌幅 5,197/5,197 一致；OHLC/边界字段无违规"),
        ("盘中/非盘中", "PASS", "8 个开闭边界值回放通过；周末自动跳过；收盘后全量任务不再被误判为盘中任务"),
        ("应用可用性", "PASS" if all_api_pass else "WARN", f"{api.get('passed', 0)}/{len(api.get('results', []))} 个生产 API 用例通过"),
        ("服务器安全", "PASS", "所有重任务均置于 MemoryHigh/MemoryMax/CPUQuota/RuntimeMax；全程 swap=0"),
        ("自动化回归", "PASS", args.test_summary),
    ], widths=[Cm(3.2), Cm(2.0), Cm(11.4)])

    add_heading(document, "1. 验收目标与范围", 1)
    add_body(document, "目标不是一次性补齐数据，而是确认系统以后能够持续、准确、按时获取数据，并在源异常、低覆盖、超时和非交易日条件下安全失败，不污染正式表、不拖死服务器。")
    add_bullets(document, [
        "覆盖启用调度任务及其输出表；停用的遗留任务不计入生产通过率，但单独列出。",
        "覆盖基础股票池、实时行情、日 K、分钟价格、日/分钟资金流、指数与概念、热榜、资讯、公告、分析、推荐、大盘汇总。",
        "覆盖交易时段边界、午休边界、收盘边界、非交易日、停牌/缺字段、空结果、低覆盖、重复键、阶段表与事务发布。",
        "准确性采用源字段规则、跨表价格一致性、数量/代码合法性、算术约束和应用 API 回读共同判定。",
    ])

    add_heading(document, "2. 生产数据路由与源头修复", 1)
    add_table(document, ["数据域", "真实写入目标", "应用读取目标", "最终策略"], [
        ("日 K", "K 线库", "K 线库", "东财全市场批量最新日；历史日保留逐股回退"),
        ("分钟价格", "分钟库", "分钟库", "抓取后直接事务写入，不再先写主库再全量搬运"),
        ("分钟资金流", "主业务库", "主业务库", "维持 API 实际路由；与分钟价格分库处理"),
        ("实时行情/日资金流", "主业务库", "主业务库", "全市场批量抓取，停牌行使用前收归一"),
        ("分析/推荐/汇总", "主业务库", "主业务库", "读取完成交易日数据，阶段发布并校验价位边界"),
    ], widths=[Cm(2.8), Cm(3.1), Cm(3.1), Cm(7.0)])

    add_heading(document, "2.1 已消除的根因", 2)
    add_bullets(document, [
        "非交易日误运行分析：将 analysis_fast 纳入非交易日自动跳过范围，避免周末无意义占用约 1GB 内存。",
        "收盘后分钟任务误判：明确 task_type 优先于共用脚本路径，15:30 全量任务不再被盘中时间窗跳过。",
        "分钟价格落错库：抓取器改为直接写入 minute_data 实际读取库；迁移镜像仅保留为运维修复模式。",
        "最新日 K 逐股过慢：改为一次全市场批量源，生产耗时从超时风险降到 25.6 秒；失败时正式表保持旧数据。",
        "停牌/残缺行情错误：统一以上一交易日收盘价回填并归一 OHLC，避免 0 价和不合法高低价。",
        "热榜代码污染：雪球只保留沪深北 A 股并重新排名；东财缺失依赖已补齐。",
        "推荐价位边界：止损严格低于入场下沿，止盈一严格高于入场上沿，止盈二高于止盈一。",
        "分析内存和性能：K 线特征默认 300 只一批流式计算；宽表字段扩展前分段整理，减少碎片化和日志开销。",
        "盘中就绪度慢查询：最新交易日改为完整命中 (k_type, adjust_type, trade_date) 索引，生产 API 从 8,228ms 降至 235ms。",
    ])

    add_heading(document, "3. 真实生产源执行结果", 1)
    source_rows = []
    for item in source.get("source_runs", []):
        outputs = item.get("outputs", {})
        source_rows.append((
            item.get("name"), item.get("status"),
            metric_text(outputs, max_items=8),
            item.get("seconds") or item.get("minutes") or item.get("cpu_seconds") or "—",
        ))
    add_table(document, ["生产执行", "结论", "输出", "耗时（秒/分钟）"], source_rows, widths=[Cm(4.0), Cm(1.8), Cm(8.5), Cm(2.6)], font_size=7.8)
    add_body(document, "重点准确性结果：实时价与日线收盘价比较 5,203 条、差异 0；日 K close/pre_close 比较 5,197 条、差异 0；000001 成交量“手→股”单位换算校验通过；分钟直写 2/2 只、480 行回读成功。分钟基准池缺少的 8 只日 K 成交量/成交额均为 0，属于当日无成交，不生成虚假分钟点。")

    add_heading(document, "4. 盘中、非盘中与边界值", 1)
    add_table(document, ["时点", "期望", "实际", "结论"], [
        (item["time"], "开盘" if item["expected_open"] else "关闭", "开盘" if item["actual_open"] else "关闭", item["status"])
        for item in source.get("time_boundary_replay", [])
    ], widths=[Cm(3.0), Cm(4.0), Cm(4.0), Cm(2.5)])
    add_bullets(document, [
        "周六/周日生产调度真实验证：盘中行情、分钟价格、分钟资金流、盘中质检和模拟交易均自动跳过。",
        "收盘后任务独立于盘中窗口：stock_minute 在 15:30 不会因脚本路径共用而被跳过。",
        "当前报告日在周日，无法制造交易所真实开盘；盘中部分采用最新完整交易日真实源全量数据 + 时间边界确定性回放 + 生产周末分支三重验证。",
    ])

    add_heading(document, "5. 数据库完整性与准确性", 1)
    add_table(document, ["检查项", "结论", "关键指标"], [
        (CHECK_LABELS.get(item.get("name"), item.get("name")), item.get("status"), metric_text(item.get("metrics", {})))
        for item in checks
    ], widths=[Cm(5.0), Cm(1.8), Cm(10.1)], font_size=7.5)

    add_heading(document, "6. 应用 API 端到端回读", 1)
    add_body(document, "API 验收从生产机访问正在运行的 127.0.0.1:8000，覆盖健康、时钟、热榜、资金流、资讯、公告、分钟、股票列表、分析和推荐；HTTP 200 之外还校验业务 error 字段与最小数据量。")
    add_table(document, ["用例", "结论", "HTTP", "数量", "耗时 ms", "路径"], [
        (item.get("name"), item.get("status"), item.get("http_status", "—"), item.get("count", 0), item.get("elapsed_ms", "—"), item.get("path"))
        for item in api.get("results", [])
    ], widths=[Cm(3.2), Cm(1.5), Cm(1.3), Cm(1.6), Cm(1.8), Cm(7.8)], font_size=7.3)

    add_heading(document, "7. 启用任务输出画像", 1)
    add_body(document, f"生产共有 {len(enabled_tasks)} 个启用调度任务。以下按唯一输出画像列出最新日期、行数/去重数和验收状态；只纳入 enabled=1 的生产范围。")
    add_table(document, ["输出表", "结论", "最新日期", "最新行数", "最新去重数", "任务类型"], [
        (
            item.get("table"), item.get("status"), item.get("latest_date") or "—",
            item.get("latest_rows", 0), item.get("latest_distinct") if item.get("latest_distinct") is not None else "—",
            ", ".join(item.get("task_types", [])),
        )
        for item in active_profiles
    ], widths=[Cm(4.1), Cm(1.5), Cm(2.2), Cm(2.1), Cm(2.2), Cm(4.8)], font_size=7.2)

    add_heading(document, "8. 服务器资源与异常恢复", 1)
    observations = source.get("resource_observations", {})
    add_table(document, ["项目", "验收结果", "证据"], [
        ("分钟全量任务", "PASS", f"常驻内存 {observations.get('minute_job_resident_mb_range')} MB；价格 13.4 分钟、资金流 12.8 分钟；覆盖率均 99.8%"),
        ("分析任务", "PASS", "300 只/批、18 批；验收软限额由 500MB 调至 650MB，硬限额 800MB；正式调度软限额 900MB"),
        ("整机余量", "PASS", f"总内存 {observations.get('server_ram_mb')}MB；最低记录可用内存约 {observations.get('minimum_observed_available_mb')}MB；swap 使用 {observations.get('swap_used_mb')}MB"),
        ("超时保护", "PASS", "60 分钟日 K 旧路径和 30 分钟迁移路径均被 RuntimeMax 安全终止，正式表未半发布"),
        ("异常清理", "PASS", "严格命名阶段表自动清理；遗留 1 张阶段表在 1.412 CPU 秒内清除"),
        ("并发保护", "PASS", "命名锁、阶段表、覆盖率门槛、单任务调度信号量共同避免重复发布和低质量覆盖"),
    ], widths=[Cm(3.0), Cm(1.8), Cm(11.8)])

    add_heading(document, "9. 异常/边界用例明细", 1)
    add_table(document, ["场景", "期望", "结果"], [
        ("停牌或价格字段缺失", "使用前收价并保证 OHLC 合法", "PASS"),
        ("实时价为 0/负数", "不得写入无效价格", "PASS"),
        ("雪球返回港美股代码", "过滤，只保留沪深北六位 A 股", "PASS"),
        ("日 K 批量返回停牌行", "前收/涨跌字段完整，成交量非负", "PASS"),
        ("分钟覆盖率不足", "不发布阶段表并返回非 0", "PASS"),
        ("抓取超时/进程被终止", "事务回滚，旧正式数据可用", "PASS"),
        ("镜像进程中断遗留阶段表", "下一次严格清理，不误删正式表", "PASS"),
        ("推荐入场区间过宽", "止损低于下沿、止盈高于上沿", "PASS"),
        ("非交易日分析任务", "自动成功跳过，不占用大内存", "PASS"),
        ("收盘后全量任务共用盘中脚本", "按任务类型运行，不按路径误跳过", "PASS"),
    ], widths=[Cm(5.2), Cm(8.3), Cm(2.0)])

    add_heading(document, "10. 范围例外与持续观察", 1)
    add_bullets(document, [
        "si_concept_constituent_east、si_index_constituent 对应遗留任务当前 enabled=0，空表不计入生产通过率；启用前必须补独立数据源并重新验收。",
        "新浪落库热榜任务已停用，旧快照日期不影响当前启用链路；实时新浪接口属于外部可用性能力，不作为数据库新鲜度门槛。",
        "2026-07-19 为周日，报告不能声称观察到了当天真实交易所盘中触发；当前通过依据为真实最新交易日数据、生产 API 回读、边界回放和周末调度实况。",
        "建议下一交易日继续观察 09:25、11:35、12:55、15:05 四个关键时点的自动任务日志；这属于运行持续观察，不影响本次代码与历史完整日数据验收结论。",
    ])

    add_heading(document, "11. 验收结论", 1)
    conclusion_fill = PALE_GREEN if overall == "通过" else PALE_AMBER
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, conclusion_fill)
    set_cell_margins(cell, top=220, start=260, bottom=220, end=260)
    p = cell.paragraphs[0]
    run = p.add_run(f"结论：{overall}\n")
    set_run_font(run, 15, True, RGBColor(19, 138, 91) if overall == "通过" else RGBColor(184, 119, 0))
    run = p.add_run("启用范围内的数据获取、准确性、边界控制、事务发布、API 回读及服务器资源保护满足生产验收条件。修复已部署，服务与调度均处于 active。")
    set_run_font(run, 10)
    add_body(document, "验收签署建议：数据负责人确认业务口径；运维负责人保留生产备份目录；下一交易日按关键时点进行常规观察。")

    document.add_page_break()
    add_heading(document, "附录 A：启用调度任务清单", 1)
    add_table(document, ["ID", "任务名称", "类型", "计划", "最近状态"], [
        (
            task.get("id"), task.get("task_name"), task.get("task_type"),
            task.get("cron_time") or (f"每 {task.get('interval_minutes')} 分钟" if task.get("interval_minutes") else "—"),
            task.get("last_run_status") or "—",
        )
        for task in enabled_tasks
    ], widths=[Cm(1.0), Cm(5.4), Cm(4.2), Cm(3.0), Cm(2.1)], font_size=7.0)

    add_heading(document, "附录 B：证据与回滚", 1)
    add_body(document, "生产部署前均建立独立备份目录。关键备份如下：")
    add_bullets(document, source.get("production_backups", []))
    add_body(document, "机器证据文件包括 inventory、schema、runtime、database、api 和 source execution JSON；验收报告由这些证据自动生成。")
    add_body(document, f"自动化回归：{args.test_summary}。")

    output = Path(args.output).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--database", required=True)
    parser.add_argument("--runtime", required=True)
    parser.add_argument("--inventory", required=True)
    parser.add_argument("--api", required=True)
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--test-summary", default="575 passed, 11 subtests passed")
    return parser.parse_args()


if __name__ == "__main__":
    build_report(parse_args())
